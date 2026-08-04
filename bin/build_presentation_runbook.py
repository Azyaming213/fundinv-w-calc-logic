"""Build the FundInv final-presentation runbook from its Markdown source."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "FINAL_PRESENTATION_RUNBOOK.md"
OUTPUT = ROOT / "docs" / "FundInv_Final_Presentation_Runbook.docx"

NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "5B6573"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "CAD2DC"
WHITE = "FFFFFF"
CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent=TABLE_INDENT_DXA) -> None:
    if sum(widths) != CONTENT_DXA:
        raise ValueError(f"Table widths must total {CONTENT_DXA} DXA")
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(CONTENT_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, (cell, width) in enumerate(zip(row.cells, widths)):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_paragraph_border(paragraph, color=BORDER, size="8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_run_font(run, name="Calibri", size=None, color=None, bold=None, italic=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


INLINE_PATTERN = re.compile(r"(\*\*.+?\*\*|`.+?`)")


def add_inline(paragraph, text: str) -> None:
    for part in INLINE_PATTERN.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, bold=True)
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=9.5, color=DARK_BLUE)
        else:
            run = paragraph.add_run(part)
            set_run_font(run)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def add_footer(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("FundInv Final Presentation Runbook  |  ")
    set_run_font(run, size=9, color=MUTED)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def add_header(section) -> None:
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run("FUNDINV  /  FINAL DEMONSTRATION")
    set_run_font(run, size=8.5, color=MUTED, bold=True)


def add_cover(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(36)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("PRESENTATION DAY OPERATOR GUIDE")
    set_run_font(run, size=10, color=BLUE, bold=True)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("FundInv Final Presentation Runbook")
    set_run_font(run, size=28, color=NAVY, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("A click-by-click, role-by-role guide for the complete client demonstration")
    set_run_font(run, size=13, color=MUTED)

    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340], indent=0)
    labels = [("12-15 min", "Live flow"), ("4 roles", "Clear duties"), ("13 funds", "Approved suite"), ("1 lifecycle", "End to end")]
    for cell, (value, label) in zip(table.rows[0].cells, labels):
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(1)
        r1 = p.add_run(value + "\n")
        set_run_font(r1, size=12, color=NAVY, bold=True)
        r2 = p.add_run(label)
        set_run_font(r2, size=9, color=MUTED)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Use this document as your speaking and recovery guide. Do not read it word-for-word.")
    set_run_font(r, size=11, color=DARK_BLUE, italic=True)
    set_paragraph_border(p, color=BLUE, size="12")


def markdown_table_widths(column_count: int) -> list[int]:
    if column_count == 2:
        return [2600, 6760]
    if column_count == 4:
        return [1500, 2700, 1700, 3460]
    base = CONTENT_DXA // column_count
    widths = [base] * column_count
    widths[-1] += CONTENT_DXA - sum(widths)
    return widths


def add_markdown_table(doc: Document, lines: list[str]) -> None:
    rows = []
    for line in lines:
        cells = [value.strip() for value in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", value) for value in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    set_table_geometry(table, markdown_table_widths(len(rows[0])))
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        for cell, value in zip(table.rows[row_index].cells, values):
            if row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            add_inline(paragraph, value)
            for run in paragraph.runs:
                set_run_font(run, size=9.5, color=NAVY if row_index == 0 else None, bold=row_index == 0)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_quote(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_BLUE)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run(" ".join(line.removeprefix("> ").strip() for line in lines))
    set_run_font(run, size=10.5, color=NAVY, italic=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT_GRAY)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("\n".join(lines))
    set_run_font(run, name="Consolas", size=9, color=NAVY)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(1)


PAGE_BREAK_HEADINGS = {
    "Recommended live demonstration flow",
    "Fund catalogue clarification",
    "Questions the audience may ask",
    "Live-demo recovery plan",
}


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    add_header(section)
    add_footer(section)
    configure_styles(doc)
    add_cover(doc)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    index = 1  # The cover already supplies the Markdown H1.
    in_code = False
    code_lines: list[str] = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(line)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        if stripped.startswith("**Audience:**") or stripped.startswith("**Recommended live-demo") or stripped.startswith("**Live portal:**"):
            index += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index])
                index += 1
            add_markdown_table(doc, table_lines)
            continue
        if stripped.startswith(">"):
            quote_lines = []
            while index < len(lines) and lines[index].strip().startswith(">"):
                quote_lines.append(lines[index].strip())
                index += 1
            add_quote(doc, quote_lines)
            continue
        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1)) - 1
            text = heading_match.group(2)
            if text in PAGE_BREAK_HEADINGS:
                doc.add_page_break()
            paragraph = doc.add_paragraph(style=f"Heading {min(level, 3)}")
            add_inline(paragraph, text)
            index += 1
            continue
        numbered = re.match(r"^\d+\.\s+(.+)$", stripped)
        if numbered:
            paragraph = doc.add_paragraph(style="List Number")
            add_inline(paragraph, numbered.group(1))
            index += 1
            continue
        bullet = re.match(r"^-\s+(.+)$", stripped)
        if bullet:
            paragraph = doc.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue

        paragraph_lines = [stripped]
        index += 1
        while index < len(lines):
            candidate = lines[index].strip()
            if not candidate:
                break
            if candidate.startswith(("#", "|", ">", "```", "- ")) or re.match(r"^\d+\.\s+", candidate):
                break
            paragraph_lines.append(candidate)
            index += 1
        paragraph = doc.add_paragraph()
        add_inline(paragraph, " ".join(paragraph_lines))

    return doc


if __name__ == "__main__":
    document = build_document()
    document.core_properties.title = "FundInv Final Presentation Runbook"
    document.core_properties.subject = "End-to-end FundInv client demonstration guide"
    document.core_properties.author = "FundInv Project Team"
    document.save(OUTPUT)
    print(OUTPUT)
