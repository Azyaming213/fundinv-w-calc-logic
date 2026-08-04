from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("FundInv_Final_Presentation_Operating_Guide.docx")

NAVY = "0D1B35"
BLUE = "2563EB"
TEAL = "0F766E"
GREEN = "15803D"
AMBER = "B45309"
RED = "B91C1C"
SLATE = "475569"
LIGHT = "F1F5F9"
PALE_BLUE = "EFF6FF"
PALE_GREEN = "F0FDF4"
PALE_AMBER = "FFFBEB"
WHITE = "FFFFFF"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margin(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, *, bold=False, color=NAVY, size=8.6) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Aptos"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margin(cell)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(SLATE)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.text = "FUNDINV  /  FINAL PRESENTATION OPERATING GUIDE"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in p.runs:
        r.font.name = "Aptos"
        r.font.size = Pt(8)
        r.font.bold = True
        r.font.color.rgb = RGBColor.from_string(BLUE)
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.add_run("Validated 4 August 2026  |  Live demo reference")
    for r in fp.runs:
        r.font.name = "Aptos"
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor.from_string(SLATE)
    fp.add_run("     ")
    add_page_number(fp)


def style_document(doc: Document) -> None:
    sec = doc.sections[0]
    sec.top_margin = Inches(0.62)
    sec.bottom_margin = Inches(0.58)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)
    add_header_footer(sec)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 12),
        ("Subtitle", 12, SLATE, 0, 10),
        ("Heading 1", 18, NAVY, 14, 8),
        ("Heading 2", 12.5, BLUE, 10, 5),
        ("Heading 3", 10.5, NAVY, 7, 3),
    ):
        style = doc.styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1"} else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_title(doc, title: str, subtitle: str | None = None) -> None:
    doc.add_heading(title, level=1)
    if subtitle:
        p = doc.add_paragraph(subtitle)
        p.style = doc.styles["Subtitle"]


def add_bullets(doc, items, *, level=0) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(3)
        if isinstance(item, tuple):
            lead, rest = item
            r = p.add_run(lead)
            r.bold = True
            p.add_run(rest)
        else:
            p.add_run(item)


def add_steps(doc, items) -> None:
    for number, item in enumerate(items, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.paragraph_format.space_after = Pt(3)
        marker = p.add_run(f"{number}.  ")
        marker.bold = True
        p.add_run(item)


def add_callout(doc, title: str, body: str, fill=PALE_BLUE, accent=BLUE) -> None:
    t = doc.add_table(rows=1, cols=2)
    t.autofit = False
    t.columns[0].width = Inches(0.09)
    t.columns[1].width = Inches(6.7)
    set_cell_width(t.cell(0, 0), 0.09)
    set_cell_width(t.cell(0, 1), 6.7)
    shade(t.cell(0, 0), accent)
    shade(t.cell(0, 1), fill)
    cell_margin(t.cell(0, 1), top=125, start=150, bottom=125, end=150)
    p = t.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Aptos"
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(accent)
    p2 = t.cell(0, 1).add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.space_before = Pt(2)
    for r2 in p2.runs:
        r2.font.name = "Aptos"
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor.from_string(NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths=None, font_size=8.0) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, color=WHITE, size=8.2)
        shade(table.rows[0].cells[i], NAVY)
        if widths:
            table.columns[i].width = Inches(widths[i])
            set_cell_width(table.rows[0].cells[i], widths[i])
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        fill = WHITE if row_idx % 2 == 0 else LIGHT
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value), size=font_size)
            shade(cells[i], fill)
            if widths:
                set_cell_width(cells[i], widths[i])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_break(doc) -> None:
    doc.add_page_break()


def build() -> None:
    doc = Document()
    style_document(doc)

    # Editorial cover
    doc.add_paragraph().paragraph_format.space_after = Pt(45)
    kicker = doc.add_paragraph("FINAL CLIENT DEMONSTRATION")
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in kicker.runs:
        r.bold = True
        r.font.name = "Aptos"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor.from_string(BLUE)
    title = doc.add_paragraph("FundInv")
    title.style = doc.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Four-role presentation operating guide")
    sub.style = doc.styles["Subtitle"]
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph().paragraph_format.space_after = Pt(18)
    add_callout(
        doc,
        "The story in one sentence",
        "FundInv separates access control, investor requests, fund valuation, and financial settlement across four roles so that no single user controls the entire transaction lifecycle.",
        fill=PALE_BLUE,
    )
    doc.add_paragraph().paragraph_format.space_after = Pt(28)
    meta = doc.add_table(rows=4, cols=2)
    meta_rows = [
        ("Audience", "OKC, assessors, and non-technical stakeholders"),
        ("Recommended demo", "12–15 minutes"),
        ("Live portal", "https://dxnh4gknox8f9.cloudfront.net"),
        ("Validation record", "Production-tested on 4 August 2026"),
    ]
    for i, (label, value) in enumerate(meta_rows):
        set_cell_text(meta.cell(i, 0), label, bold=True, color=BLUE, size=9)
        set_cell_text(meta.cell(i, 1), value, size=9)
        shade(meta.cell(i, 0), LIGHT)
        shade(meta.cell(i, 1), WHITE)
    p = doc.add_paragraph("Use this document as the presenter’s script and safety checklist. It records what was actually verified, what remains externally blocked, and what each control means.")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)

    page_break(doc)
    add_title(doc, "1. Readiness verdict", "What is safe to demonstrate tomorrow")
    add_callout(
        doc,
        "Core workflow: demonstration-ready",
        "Admin access, Investor fund browsing, demo PayNow subscription, Operations settlement, Alpaca-assisted Manager valuation with preview, Investor P&L allocation, redemption, holdings, analytics, audit records, and AWS-hosted access were exercised successfully.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_table(
        doc,
        ["Area", "Result", "Evidence"],
        [
            ("Backend test suite", "PASS", "36/36 unit and integration tests"),
            ("Frontend production build", "PASS", "36 routes compiled and prerendered"),
            ("Four-role browser audit", "PASS", "10/10 deployed route, access, and cross-role checks"),
            ("Live subscription", "PASS", "$12.34 matched, settled once"),
            ("Live redemption", "PASS", "$7.89 approved, transferred, completed"),
            ("P&L allocation", "PASS", "$10 fund P&L allocated $4.00 to a 40% holder"),
            ("Investor dashboard consistency", "PASS", "Value changed by net cash flow of +$4.45"),
            ("Automatic market suggestion", "PASS", "AWS Alpaca account/data calls return 200; VOO prefill verified"),
            ("CI/CD credentials", "ATTENTION", "GitHub deployment jobs fail at AWS credential setup; manual rollout used"),
        ],
        widths=[1.55, 0.8, 4.3],
    )
    add_callout(
        doc,
        "Do not claim zero defects",
        "The demonstrated workflows and controls have passed the current test suite, but no finite test proves zero defects or enterprise readiness. Present the system as a tested capstone deployment with documented operational limits.",
        fill=PALE_AMBER,
        accent=AMBER,
    )

    page_break(doc)
    add_title(doc, "2. The four-role control model", "Who does what, and why")
    add_table(
        doc,
        ["Role", "Owns", "Must not do", "Hands off to"],
        [
            ("Admin", "Invitations, users, role/access status, oversight", "Value funds or settle money", "Investor / Manager / Operations"),
            ("Investor", "Chooses a fund and submits subscription/redemption requests", "Issue units, set NAV, self-approve", "Operations"),
            ("Manager", "Reviews fund performance, previews and finalizes NAV/P&L", "Confirm payment or alter investor requests", "Operations and Investor reporting"),
            ("Operations", "Verifies money movement and settles units", "Change P&L, NAV, or user roles", "Investor reporting and Admin audit"),
        ],
        widths=[0.85, 2.0, 1.8, 1.7],
    )
    doc.add_heading("End-to-end sequence", level=2)
    add_steps(doc, [
        "Admin grants the correct role and leaves an auditable access record.",
        "Investor selects a fund, account, and exact amount; demo PayNow records the same amount and reference.",
        "Manager finalizes the applicable daily fund valuation so a valid NAV exists.",
        "Operations compares request, receipt, fund, status, and valuation before settlement.",
        "The database issues or redeems internal fund units exactly once.",
        "Investor sees updated holdings, activity, NAV P&L allocation, and analytics; Admin can audit the lifecycle.",
    ])
    add_callout(
        doc,
        "Presentation language",
        "Say fund units, subscription, redemption, finalized NAV, and demo PayNow. Do not say the Investor directly owns public QQQ or VOO shares; the Investor owns internal units in a FundInv product.",
    )

    page_break(doc)
    add_title(doc, "3. Manager daily valuation — the straight explanation")
    add_callout(
        doc,
        "Yes: the field is intended to prefill automatically",
        "Refresh calculation should obtain Alpaca market prices, calculate a weighted fund return, multiply it by opening assets, and prefill Suggested daily P&L. The Manager reviews that suggestion rather than searching the internet for a dollar P&L figure.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    doc.add_heading("Where the number comes from", level=2)
    add_table(
        doc,
        ["Input", "Source", "Purpose"],
        [
            ("Latest/current price", "Alpaca snapshot or latest daily bar", "End-of-period market value proxy"),
            ("Previous close", "Alpaca prior daily bar", "Return comparison point"),
            ("Component weight", "Fund configuration in PostgreSQL", "Combines multi-component returns"),
            ("Opening assets", "Previous finalized valuation", "Dollar base for suggested P&L"),
            ("Opening units", "Settled fund-unit ledger", "Denominator for NAV"),
        ],
        widths=[1.35, 2.5, 2.5],
    )
    doc.add_heading("Calculation", level=2)
    add_bullets(doc, [
        ("Component return = ", "current price / previous close − 1"),
        ("Fund return = ", "sum of component weight × component return"),
        ("Suggested daily P&L = ", "opening assets × fund return"),
        ("Closing assets before flows = ", "opening assets + daily P&L"),
        ("NAV per unit = ", "closing assets before flows / opening units"),
    ])
    add_callout(
        doc,
        "Current production condition",
        "AWS now authenticates successfully to Alpaca paper trading and IEX market data. The deployed VOO check for 3 August 2026 returned a +1.45681278% market return and automatically suggested +$74.3703 on $5,105.00 opening assets. The Manager still reviews the prices, date, weights and preview before finalizing.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    doc.add_heading("What the Manager clicks", level=2)
    add_steps(doc, [
        "Choose a managed fund and a date not already present in Finalized valuation history.",
        "Click Refresh calculation. If a suggestion appears, review its return, prices, weights, and timestamp.",
        "Use the prefilled suggestion when the source, date and component figures are correct. Override it only for an approved accounting adjustment and explain that adjustment in Audit note.",
        "Click Preview calculation. Check opening assets, P&L, closing assets, units, NAV, and investor allocations.",
        "Finalize once only when the preview is correct. The record becomes historical and is not casually editable.",
    ])

    page_break(doc)
    add_title(doc, "4. Verified P&L example", "VOO on 4 August 2026")
    add_table(
        doc,
        ["Metric", "Verified value", "Meaning"],
        [
            ("Opening assets", "$5,105.00", "Fund assets before the day’s P&L and flows"),
            ("Manager-approved daily P&L", "+$10.00", "Manual presentation-readiness value; not Alpaca-derived"),
            ("Closing assets before flows", "$5,115.00", "$5,105.00 + $10.00"),
            ("Opening units", "5,000.0000", "Units outstanding before same-day flows"),
            ("Finalized NAV", "$1.02300000", "$5,115.00 / 5,000.0000"),
            ("Test Investor ownership", "40.00%", "2,000 / 5,000 opening units"),
            ("Test Investor allocated P&L", "+$4.00", "40% × $10.00"),
            ("QA Investor allocated P&L", "+$0.78", "7.8355% × $10.00, displayed rounded"),
            ("External ownership allocation", "+$5.22", "Remainder; all allocations sum to $10.00"),
        ],
        widths=[2.05, 1.45, 2.85],
    )
    add_callout(
        doc,
        "What performance changes",
        "Daily P&L changes NAV and each holder’s value. It does not create units. Subscriptions create units and redemptions remove units at the finalized NAV; those cash flows are kept separate from investment performance.",
    )
    doc.add_heading("Verified same-day settlement arithmetic", level=2)
    add_table(
        doc,
        ["Action", "Amount", "Units at $1.023 NAV", "Result"],
        [
            ("VOO subscription", "$12.34", "+12.0626", "Completed by Operations once"),
            ("VOO redemption", "$7.89", "−7.7126", "Approved, payout stage, completed"),
            ("Net", "+$4.45", "+4.3500", "2,000.0000 → 2,004.3500 units"),
        ],
        widths=[1.6, 1.0, 1.55, 2.2],
    )
    p = doc.add_paragraph("Portfolio value moved from $5,272.00 to $5,276.45, exactly matching the net cash flow of +$4.45. Daily P&L remained $4.00 because cash flows do not masquerade as performance.")
    p.runs[0].bold = True

    page_break(doc)
    add_title(doc, "5. Admin operator reference")
    add_table(
        doc,
        ["Page / control", "What it does", "Role handoff / presentation point"],
        [
            ("Dashboard", "Shows access and oversight summary", "Start here; Admin observes rather than settles"),
            ("Users → Edit", "Changes full name, email, role, password reset value, active status", "Confirm role and status; important changes are audited"),
            ("Users → Active", "Deactivates access without erasing financial history", "Prefer deactivation; no destructive delete-user control"),
            ("Settings → Invite User", "Fields: full name, email, role; creates/sends invitation", "Hands a new identity to the selected role workflow"),
            ("Pending invitations → Resend", "Resends an existing invitation", "Use only for a real prepared address"),
            ("Pending invitations → Delete", "Removes the invitation, not a financial user ledger", "Do not confuse invitation deletion with user deletion"),
            ("Operations invite requests → Approve / Reject", "Reviews access requests raised by Operations", "Shows two-person access administration"),
            ("Investor-manager assignment", "Maps an Investor to a Manager", "Determines which Manager sees/serves that Investor"),
            ("Fund Flows", "Read-only financial oversight", "Operations owns settlement"),
            ("Valuations", "Reviews finalized NAV/P&L history", "Manager owns finalization"),
            ("Transactions", "Searches and filters recorded activity", "Evidence and reconciliation"),
            ("Audit Logs", "Shows actor, role, action, time, and entity details", "Close the demo with traceability"),
            ("Articles", "Creates/edits educational content", "Visible to Investors without changing holdings"),
            ("Account security", "2FA and account-security controls", "Set up before the presentation; do not spend live time here"),
        ],
        widths=[1.65, 2.5, 2.5],
        font_size=7.6,
    )
    add_callout(doc, "Adding a user proof", "A QA Investor account already exists and has a settled VOO holding. For the live demo, create or show an invitation using a prepared address; do not send an unplanned email to a real person.")

    page_break(doc)
    add_title(doc, "6. Investor operator reference")
    add_table(
        doc,
        ["Control", "What it actually does", "Downstream role / safeguard"],
        [
            ("Account selector", "Chooses the portfolio account being viewed or used", "Flows and holdings remain linked to that account"),
            ("Add Account", "Creates another portfolio account/strategy container", "Does not add cash or units"),
            ("Edit strategy", "Updates the account strategy label", "Does not rebalance broker holdings automatically"),
            ("Deposit to Fund", "Select account, fund, exact amount; creates subscription", "Demo PayNow must record the same amount; Operations settles"),
            ("Demo QR / Simulate payment", "Displays a fixed dummy QR and records exact demo receipt", "No real money moves; no units yet"),
            ("Withdraw", "Select account, settled fund position, amount", "Cannot exceed settled redeemable value; Operations settles"),
            ("Funds", "Search, sort, and filter available fund products", "Invest opens the subscription workflow"),
            ("Fund Flows", "Shows request ID, type, amount, payment state, and status", "Explains whether Investor, Manager, or Operations is blocking progress"),
            ("NAV P&L Allocations", "Shows finalized daily allocation for this Investor", "Derived from Manager-finalized fund P&L and opening ownership"),
            ("Account Value", "Settled holdings valued at latest finalized NAV", "Pending requests do not become holdings"),
            ("Daily / monthly / YTD", "Shows performance metrics separated from cash flows", "Deposits/withdrawals must not appear as investment return"),
            ("Performance dates + Calculate", "Recalculates the selected analytics period", "Use a populated period during the demo"),
            ("Fund Allocation", "Pie chart of settled value by fund product", "Uses the same authoritative positions as the table"),
            ("Recent Fund Activity", "Latest subscriptions/redemptions", "Replaced unrelated AAPL/MSFT sample transactions"),
            ("Fund Investments", "All settled positions: units, NAV, cost basis, value", "Now matches the allocation chart"),
            ("PDF", "Exports portfolio summary", "Demonstrate only after holdings load"),
            ("Email Summary", "Emails portfolio summary to the signed-in Investor", "External SMTP delivery may be delayed"),
            ("Articles", "Reads educational content by category", "No effect on financial records"),
            ("Feedback", "Fields: subject and message; submits service feedback", "Operations can respond"),
            ("Account security", "2FA/account-security controls", "Preconfigure before the live run"),
        ],
        widths=[1.45, 2.55, 2.65],
        font_size=7.35,
    )

    page_break(doc)
    add_title(doc, "7. Manager operator reference")
    add_table(
        doc,
        ["Page / control", "What it actually does", "Relationship to other roles"],
        [
            ("Dashboard / Investors", "Shows Investors assigned to this Manager", "Admin controls assignments"),
            ("Funds → Your Funds", "Shows all 13 assigned demo products plus future Manager-created products", "Operations approval makes new funds investable"),
            ("Create Fund: name / description", "Defines the product identity", "Does not create an Alpaca fund or investor position"),
            ("Strategy / risk", "Classifies mandate and risk presentation", "Investor sees these catalogue labels"),
            ("Underlying search", "Adds ticker components from market lookup", "Defines the simplified market-return model"),
            ("Allocation weights / Remove", "Sets component weights; total should be 100%", "Weights drive automatic suggested return"),
            ("Create", "Submits a fund for Operations review", "It becomes visible/investable only after approval"),
            ("Performance dates / Refresh", "Loads historical performance attribution", "Manager analysis; does not modify holdings"),
            ("Weight fields / Run what-if", "Models a hypothetical allocation", "Scenario only; does not trade or finalize NAV"),
            ("Valuation Fund / Date", "Chooses one managed fund and business date", "Date must not already be finalized"),
            ("Refresh calculation", "Requests and prefills the Alpaca market-data suggestion", "Manager reviews it before preview/finalization"),
            ("Daily fund P&L", "Prefilled suggestion or externally verified manual fallback", "Changes NAV/value, never unit count"),
            ("Audit note", "Records source/reason for manual value or adjustment", "Admin/Operations audit evidence"),
            ("Preview calculation", "Shows assets, units, NAV and ownership allocation before committing", "Required review point before Operations settles"),
            ("Finalize valuation", "Writes one finalized daily valuation", "Unlocks applicable Operations settlement"),
            ("Finalized history", "Shows past P&L, NAV, date, source, actor", "Evidence for Investor allocations and audit"),
            ("Transactions / search / side / CSV", "Inspects activity and exports reconciliation data", "Read/analysis function, not settlement"),
            ("Articles / security", "Educational content and account protection", "No financial authority"),
        ],
        widths=[1.65, 2.55, 2.45],
        font_size=7.35,
    )

    page_break(doc)
    add_title(doc, "8. Operations operator reference")
    add_table(
        doc,
        ["Control", "What it actually does", "Required check / handoff"],
        [
            ("Fund Flows search", "Searches email, name, or request ID", "Use the Investor’s unique reference"),
            ("Type / status filters", "Narrows subscription, redemption, and lifecycle state", "Use Pending Ops Team during the demo"),
            ("Requested / Received", "Shows fixed requested amount and recorded receipt", "Must match before a demo subscription settles"),
            ("Verify & Complete", "One-step demo PayNow subscription settlement", "Issues units once at finalized NAV; action disappears after completion"),
            ("Approve redemption", "Authorises the request and moves it toward outgoing transfer", "Does not yet remove units"),
            ("Complete redemption", "Confirms payout was sent and redeems units", "Use only after outgoing transfer; action is one-time"),
            ("Start / Retry Payout", "Provider payout action when Stripe Connect is used", "Not needed for manual/demo path"),
            ("Reject", "Rejects a pending request", "A reason should explain the decision in notes/audit"),
            ("Settlement status message", "Explains missing valuation/payment/payout condition", "Read it instead of repeatedly clicking"),
            ("Fund Reviews → Approve / Reject", "Reviews Manager-created fund products", "Approval makes eligible products investable"),
            ("Audit Logs", "Reviews sensitive actions", "Confirms actor, role, time, entity"),
            ("Feedback → Respond", "Closes Investor feedback loop", "Service task, not financial settlement"),
            ("Invite Requests", "Fields: full name, email, requested role", "Admin independently approves/rejects"),
            ("Account security", "2FA/account protection", "Set up before presenting"),
        ],
        widths=[1.55, 2.6, 2.5],
        font_size=7.45,
    )
    add_callout(
        doc,
        "Why withdrawal has two Operations actions",
        "Approval confirms the redemption is authorised. Completion confirms the outgoing payout has actually been sent and only then removes units. Demo PayNow subscriptions use one Verify & Complete action because receipt is already recorded before Operations review.",
    )

    page_break(doc)
    add_title(doc, "9. Product suite and prepared Investor")
    doc.add_heading("Catalogue truth", level=2)
    add_bullets(doc, [
        "The original seed had 20 rows: 13 fund products and 7 ordinary stocks.",
        "The seven stock rows were removed from the Investor fund catalogue.",
        "The valid demonstration suite is 13 products: QQQ, VOO, VTI, SPY, BND, AGG, VYM, SCHD, SOXL, TQQQ, TLT, SCHR, and VIG.",
        "Manager ‘Your Funds (3)’ is a manager-ownership count, not the total catalogue count.",
    ])
    doc.add_heading("investor@fundinv.com now owns 10 settled products", level=2)
    add_table(
        doc,
        ["Established", "Added through normal demo workflow"],
        [
            ("QQQ — Invesco QQQ Trust", "TLT — iShares 20+ Year Treasury Bond"),
            ("VOO — Vanguard S&P 500 ETF", "AGG — iShares Core US Aggregate Bond"),
            ("", "SCHR — Schwab Short-Term US Treasury ETF"),
            ("", "SCHD — Schwab US Dividend Equity ETF"),
            ("", "VIG — Vanguard Dividend Appreciation ETF"),
            ("", "VYM — Vanguard High Dividend Yield"),
            ("", "BND — Vanguard Total Bond Market"),
            ("", "VTI — Vanguard Total Stock Market"),
        ],
        widths=[3.05, 3.3],
    )
    add_callout(
        doc,
        "How this was achieved",
        "Eight $25 subscriptions were created through the Investor fund catalogue, paid with the demo PayNow action, and completed by Operations. No direct database balance edit was used. SPY has a paid request but remains pending because its Manager valuation/mapping condition is not satisfied; SOXL and TQQQ are not part of this Investor’s prepared holdings.",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    page_break(doc)
    add_title(doc, "10. Presentation run of show", "12–15 minutes, business-first")
    add_table(
        doc,
        ["Time", "Role / screen", "Action", "What to say"],
        [
            ("0:00", "Opening", "State the four-role control model", "Access, valuation, settlement, and reporting are separated"),
            ("0:45", "Admin", "Show prepared invitation or create one", "Admin grants roles but cannot settle money"),
            ("2:00", "Investor / Funds", "Search VOO; open Invest", "Investor chooses a fund product and exact amount"),
            ("3:00", "Investor / QR", "Submit and simulate demo PayNow", "Amount/reference are locked; no units yet"),
            ("5:00", "Manager / Valuation", "Refresh, explain source, preview, show finalized history", "NAV changes unit value, not unit count"),
            ("8:00", "Operations / Flows", "Match request and receipt; Verify & Complete", "Independent settlement; one-time action"),
            ("10:00", "Investor / Portfolio", "Show activity, positions, P&L, chart", "Cash flow and performance remain separate"),
            ("11:30", "Investor / Withdraw", "Submit safe redemption", "Cannot exceed settled position"),
            ("12:30", "Operations", "Approve then complete after payout step", "Two stages protect outgoing money"),
            ("14:00", "Admin / Audit", "Show audit evidence and close", "Every sensitive step has accountable ownership"),
        ],
        widths=[0.55, 1.45, 2.15, 2.55],
        font_size=7.5,
    )
    doc.add_heading("Pre-login setup", level=2)
    add_bullets(doc, [
        "Use four separate browser profiles or a mix of normal/private windows. Tabs in one browser profile share the same authentication cookie and will overwrite each other’s role session.",
        "Admin: admin@fundinv.com / admin123",
        "Manager: manager@fundinv.com / admin123",
        "Operations: operations@fundinv.com / admin123",
        "Investor: investor@fundinv.com / investor123",
        "Use localhost only for local testing; use the CloudFront URL for the hosted presentation.",
    ])

    page_break(doc)
    add_title(doc, "11. Seamlessness and remaining limitations")
    add_callout(
        doc,
        "Overall assessment: strong classroom demo, not defect-free production accounting",
        "The role handoffs and core accounting controls are coherent and usable. The demo is smooth when role windows and an unfinalized valuation date are prepared. Alpaca market-data authentication now works; CI/CD credentials and single-instance availability remain operational limitations.",
        fill=PALE_AMBER,
        accent=AMBER,
    )
    add_table(
        doc,
        ["Issue", "Presentation impact", "Recommended handling"],
        [
            ("External market-data outage", "Automatic suggestion may be temporarily unavailable", "Use a prepared authoritative figure with an audit note; do not invent or Google dollar P&L"),
            ("GitHub Actions AWS credentials expired/invalid", "Push does not auto-deploy", "Current release was manually deployed; repair OIDC/static credential setup after presentation"),
            ("Single-instance ASG desired capacity", "Good for small demo, not highly available", "Do not claim enterprise HA; load test before scale claims"),
            ("SPY paid request pending", "Visible blocker if selected", "Use VOO or QQQ for the live workflow"),
            ("Manual manager_entry provenance", "System trusts external Manager source", "Use audit note and authoritative statement; never Google dollar P&L"),
            ("Simplified market-return valuation", "Does not include exact broker quantities, cash, fees, dividends, liabilities", "Describe as a prototype valuation model with Manager accounting adjustment"),
            ("Shared Alpaca holdings concept", "Could be mistaken for Investor-specific holdings", "Underlying holdings section is hidden/redesigned; use Fund Investments for Investor ownership"),
        ],
        widths=[1.8, 2.2, 2.4],
        font_size=7.65,
    )

    page_break(doc)
    add_title(doc, "12. Final preflight and recovery checklist")
    doc.add_heading("Ten minutes before presenting", level=2)
    add_bullets(doc, [
        "Open the CloudFront portal in four isolated browser profiles and confirm each role is still signed in.",
        "Keep VOO or QQQ selected; do not use SPY for the transaction demo.",
        "Confirm the valuation date is not already finalized before promising a live finalization.",
        "Prepare a small, memorable subscription and redemption amount.",
        "Keep paynow_demo; do not use real money or Stripe payout actions.",
        "Have this guide open on a second screen and keep credentials out of the projected browser history.",
        "Do not edit AWS secrets, email settings, roles, or fund mappings during the presentation.",
    ])
    doc.add_heading("If something stops", level=2)
    add_table(
        doc,
        ["Symptom", "Meaning", "Safe response"],
        [
            ("Date already finalized", "Final records are protected", "Choose a truly unfinalized date or show history"),
            ("Market suggestion unavailable", "External market service failed/auth failed", "Use prepared verified figure with audit note; state limitation"),
            ("Operations button disabled", "Payment, valuation, fund, or state prerequisite is missing", "Read the status message; do not repeatedly click"),
            ("Investor values look stale", "Previous role mutation has not been reloaded", "Refresh only after Manager/Operations completes"),
            ("Invitation email delayed", "Notification provider delay", "Show invitation record; access control remains in the database"),
            ("Role login replaces another tab", "Same browser cookie shared", "Use separate profiles/private windows"),
        ],
        widths=[1.45, 2.15, 2.8],
        font_size=7.8,
    )
    add_callout(
        doc,
        "Closing line",
        "FundInv demonstrates a complete controlled lifecycle: Admin grants access, the Investor requests a transaction, the Manager values the fund, Operations independently settles money and units, and the Investor receives transparent holdings and performance reporting.",
        fill=PALE_GREEN,
        accent=GREEN,
    )

    doc.core_properties.title = "FundInv Final Presentation Operating Guide"
    doc.core_properties.subject = "Four-role end-to-end demonstration runbook"
    doc.core_properties.author = "FundInv"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
